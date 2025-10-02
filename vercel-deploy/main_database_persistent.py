"""
Document Intelligence Platform - Database-Persistent Vercel Version
Uses PostgreSQL for persistent storage of documents and embeddings
"""
import os
import json
import time
import hashlib
import tempfile
import io
from typing import List, Dict, Any, Optional
from pathlib import Path

from fastapi import FastAPI, HTTPException, UploadFile, File, Form, Depends
from fastapi.responses import HTMLResponse, RedirectResponse, JSONResponse
from pydantic import BaseModel, Field
from pydantic_settings import BaseSettings
import psycopg2
from psycopg2.extras import RealDictCursor
import numpy as np
from openai import OpenAI
import PyPDF2
from docx import Document

# Configuration
class Settings(BaseSettings):
    openai_api_key: str = Field(default="", alias="OPENAI_API_KEY")
    database_url: str = Field(default="", alias="DATABASE_URL")
    app_name: str = Field(default="Document Intelligence Platform", alias="APP_NAME")
    app_version: str = Field(default="1.0.0", alias="APP_VERSION")
    
    class Config:
        env_file = ".env"
        case_sensitive = False
        extra = "ignore"

settings = Settings()

# Initialize FastAPI app
app = FastAPI(
    title=settings.app_name,
    description="AI-powered document analysis with RAG capabilities",
    version=settings.app_version
)

# Initialize OpenAI client
client = OpenAI(api_key=settings.openai_api_key) if settings.openai_api_key else None

# Database connection
def get_db_connection():
    """Get database connection"""
    try:
        if settings.database_url:
            conn = psycopg2.connect(settings.database_url)
            return conn
        else:
            print("No database URL configured")
            return None
    except Exception as e:
        print(f"Database connection error: {str(e)}")
        return None

def init_database():
    """Initialize database tables"""
    conn = get_db_connection()
    if not conn:
        return False
    
    try:
        cur = conn.cursor()
        
        # Create documents table
        cur.execute("""
            CREATE TABLE IF NOT EXISTS documents (
                id SERIAL PRIMARY KEY,
                doc_id VARCHAR(50) UNIQUE NOT NULL,
                filename VARCHAR(255) NOT NULL,
                file_type VARCHAR(100) NOT NULL,
                file_size INTEGER NOT NULL,
                status VARCHAR(50) DEFAULT 'processed',
                chunks_count INTEGER DEFAULT 0,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                metadata JSONB
            )
        """)
        
        # Create chunks table
        cur.execute("""
            CREATE TABLE IF NOT EXISTS document_chunks (
                id SERIAL PRIMARY KEY,
                chunk_id VARCHAR(100) UNIQUE NOT NULL,
                document_id VARCHAR(50) NOT NULL,
                chunk_index INTEGER NOT NULL,
                content TEXT NOT NULL,
                embedding VECTOR(1536),
                metadata JSONB,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """)
        
        # Create queries table
        cur.execute("""
            CREATE TABLE IF NOT EXISTS queries (
                id SERIAL PRIMARY KEY,
                query_id VARCHAR(50) UNIQUE NOT NULL,
                query_text TEXT NOT NULL,
                response TEXT,
                processing_time FLOAT,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """)
        
        conn.commit()
        cur.close()
        conn.close()
        print("Database initialized successfully")
        return True
        
    except Exception as e:
        print(f"Database initialization error: {str(e)}")
        conn.rollback()
        conn.close()
        return False

# Initialize database on startup
init_database()

# Pydantic models
class QueryRequest(BaseModel):
    query: str
    k: int = Field(default=5, ge=1, le=20)
    use_cache: bool = True
    rerank_context: bool = True

class DocumentChunk(BaseModel):
    id: str
    content: str
    metadata: Dict[str, Any]
    score: float

class QueryResponse(BaseModel):
    query_id: str
    query: str
    response: str
    context_documents: List[str]
    context_chunks: List[DocumentChunk]
    similarity_scores: List[float]
    processing_time: float
    cached: bool = False

class DocumentInfo(BaseModel):
    id: str
    filename: str
    file_type: str
    file_size: int
    status: str
    chunks_count: int
    created_at: str
    metadata: Dict[str, Any]

class StatsResponse(BaseModel):
    total_documents: int
    total_chunks: int
    total_queries: int
    avg_processing_time: float

# Document processing functions
def extract_text_from_pdf(file_content: bytes) -> str:
    """Extract text from PDF file"""
    try:
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error reading PDF: {str(e)}")

def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from DOCX file"""
    try:
        doc = Document(io.BytesIO(file_content))
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error reading DOCX: {str(e)}")

def extract_text_from_txt(file_content: bytes) -> str:
    """Extract text from TXT file"""
    try:
        return file_content.decode('utf-8')
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error reading TXT: {str(e)}")

def split_text_into_chunks(text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
    """Split text into overlapping chunks"""
    chunks = []
    start = 0
    
    while start < len(text):
        end = start + chunk_size
        chunk = text[start:end]
        
        # Try to break at sentence boundary
        if end < len(text):
            last_period = chunk.rfind('.')
            last_newline = chunk.rfind('\n')
            break_point = max(last_period, last_newline)
            
            if break_point > start + chunk_size // 2:
                chunk = text[start:start + break_point + 1]
                end = start + break_point + 1
        
        chunks.append(chunk.strip())
        start = end - overlap
    
    return [chunk for chunk in chunks if chunk.strip()]

def generate_embedding(text: str) -> List[float]:
    """Generate embedding using OpenAI"""
    if not client:
        print("OpenAI client not available - returning dummy embedding")
        return [0.1] * 1536
    
    try:
        response = client.embeddings.create(
            model="text-embedding-3-small",
            input=text
        )
        embedding = response.data[0].embedding
        print(f"Generated embedding with length {len(embedding)}")
        return embedding
    except Exception as e:
        print(f"Error generating embedding: {str(e)}")
        return [0.1] * 1536

def cosine_similarity(a: List[float], b: List[float]) -> float:
    """Calculate cosine similarity between two vectors"""
    dot_product = sum(x * y for x, y in zip(a, b))
    norm_a = sum(x * x for x in a) ** 0.5
    norm_b = sum(y * y for y in b) ** 0.5
    
    if norm_a == 0 or norm_b == 0:
        return 0.0
    
    return dot_product / (norm_a * norm_b)

def process_document(file_content: bytes, filename: str, file_type: str) -> Dict[str, Any]:
    """Process uploaded document and extract text"""
    conn = get_db_connection()
    if not conn:
        raise HTTPException(status_code=500, detail="Database connection failed")
    
    try:
        cur = conn.cursor()
        
        # Extract text based on file type
        if file_type == "application/pdf":
            text = extract_text_from_pdf(file_content)
        elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            text = extract_text_from_docx(file_content)
        elif file_type == "text/plain":
            text = extract_text_from_txt(file_content)
        else:
            raise HTTPException(status_code=400, detail=f"Unsupported file type: {file_type}")
        
        if not text.strip():
            raise HTTPException(status_code=400, detail="No text content found in document")
        
        # Split text into chunks
        chunks = split_text_into_chunks(text)
        
        # Generate document ID
        doc_id = f"doc_{int(time.time())}"
        
        # Insert document record
        cur.execute("""
            INSERT INTO documents (doc_id, filename, file_type, file_size, chunks_count, metadata)
            VALUES (%s, %s, %s, %s, %s, %s)
            ON CONFLICT (doc_id) DO UPDATE SET
                filename = EXCLUDED.filename,
                file_size = EXCLUDED.file_size,
                chunks_count = EXCLUDED.chunks_count,
                metadata = EXCLUDED.metadata
        """, (
            doc_id,
            filename,
            file_type,
            len(file_content),
            len(chunks),
            json.dumps({
                "total_chunks": len(chunks),
                "total_characters": len(text),
                "processing_time": time.time()
            })
        ))
        
        # Process chunks
        for i, chunk in enumerate(chunks):
            chunk_id = f"{doc_id}_chunk_{i}"
            embedding = generate_embedding(chunk)
            
            # Insert chunk record
            cur.execute("""
                INSERT INTO document_chunks (chunk_id, document_id, chunk_index, content, embedding, metadata)
                VALUES (%s, %s, %s, %s, %s, %s)
                ON CONFLICT (chunk_id) DO UPDATE SET
                    content = EXCLUDED.content,
                    embedding = EXCLUDED.embedding,
                    metadata = EXCLUDED.metadata
            """, (
                chunk_id,
                doc_id,
                i,
                chunk,
                json.dumps(embedding),  # Store as JSON for now
                json.dumps({
                    "filename": filename,
                    "chunk_index": i
                })
            ))
        
        conn.commit()
        
        # Get the created document
        cur.execute("SELECT * FROM documents WHERE doc_id = %s", (doc_id,))
        doc_record = cur.fetchone()
        
        cur.close()
        conn.close()
        
        return {
            "id": doc_record[1],  # doc_id
            "filename": doc_record[2],
            "file_type": doc_record[3],
            "file_size": doc_record[4],
            "status": doc_record[5],
            "chunks_count": doc_record[6],
            "created_at": doc_record[7].isoformat() if doc_record[7] else "",
            "metadata": doc_record[8] if doc_record[8] else {}
        }
        
    except Exception as e:
        conn.rollback()
        cur.close()
        conn.close()
        raise HTTPException(status_code=500, detail=str(e))

def search_similar_chunks(query: str, k: int = 5) -> List[DocumentChunk]:
    """Search for similar chunks using database"""
    conn = get_db_connection()
    if not conn:
        print("No database connection - trying text-based search")
        return text_based_search(query, k)
    
    try:
        cur = conn.cursor(cursor_factory=RealDictCursor)
        
        # Generate query embedding
        query_embedding = generate_embedding(query)
        
        if not query_embedding or all(x == 0.0 for x in query_embedding):
            print("Query embedding generation failed - trying text-based search")
            return text_based_search(query, k)
        
        print(f"Searching database for query: {query[:50]}...")
        
        # Get all chunks with embeddings
        cur.execute("""
            SELECT chunk_id, document_id, chunk_index, content, embedding, metadata
            FROM document_chunks 
            WHERE embedding IS NOT NULL
        """)
        
        chunks = cur.fetchall()
        print(f"Found {len(chunks)} chunks in database")
        
        if not chunks:
            print("No chunks found in database")
            return []
        
        # Calculate similarities
        similarities = []
        for chunk in chunks:
            try:
                chunk_embedding = json.loads(chunk['embedding'])
                similarity = cosine_similarity(query_embedding, chunk_embedding)
                similarities.append((chunk, similarity))
            except Exception as e:
                print(f"Error processing chunk {chunk['chunk_id']}: {str(e)}")
                continue
        
        if not similarities:
            print("No valid similarities found - trying text-based search")
            return text_based_search(query, k)
        
        # Sort by similarity and get top k
        similarities.sort(key=lambda x: x[1], reverse=True)
        top_chunks = similarities[:k]
        
        print(f"Top similarities: {[(chunk['chunk_id'], f'{score:.3f}') for chunk, score in top_chunks[:3]]}")
        
        results = []
        for chunk, score in top_chunks:
            metadata = json.loads(chunk['metadata']) if chunk['metadata'] else {}
            results.append(DocumentChunk(
                id=chunk['chunk_id'],
                content=chunk['content'],
                metadata={
                    "document_id": chunk['document_id'],
                    "filename": metadata.get("filename", "unknown"),
                    "chunk_index": chunk['chunk_index']
                },
                score=score
            ))
        
        cur.close()
        conn.close()
        
        print(f"Returning {len(results)} chunks from database search")
        return results
        
    except Exception as e:
        print(f"Error in database search: {str(e)} - trying text-based search")
        if conn:
            conn.close()
        return text_based_search(query, k)

def text_based_search(query: str, k: int = 5) -> List[DocumentChunk]:
    """Fallback text-based search using database"""
    print(f"Performing text-based search for: {query}")
    
    conn = get_db_connection()
    if not conn:
        print("No database connection for text search")
        return []
    
    try:
        cur = conn.cursor(cursor_factory=RealDictCursor)
        
        # Search for chunks containing query words
        query_words = query.lower().split()
        search_pattern = '%' + '%'.join(query_words) + '%'
        
        cur.execute("""
            SELECT chunk_id, document_id, chunk_index, content, metadata
            FROM document_chunks 
            WHERE LOWER(content) LIKE %s
            ORDER BY LENGTH(content) ASC
            LIMIT %s
        """, (search_pattern, k))
        
        chunks = cur.fetchall()
        print(f"Text search found {len(chunks)} chunks")
        
        results = []
        for chunk in chunks:
            metadata = json.loads(chunk['metadata']) if chunk['metadata'] else {}
            results.append(DocumentChunk(
                id=chunk['chunk_id'],
                content=chunk['content'],
                metadata={
                    "document_id": chunk['document_id'],
                    "filename": metadata.get("filename", "unknown"),
                    "chunk_index": chunk['chunk_index']
                },
                score=0.5  # Default score for text search
            ))
        
        cur.close()
        conn.close()
        
        return results
        
    except Exception as e:
        print(f"Error in text-based search: {str(e)}")
        if conn:
            conn.close()
        return []

def generate_response(query: str, context_chunks: List[DocumentChunk]) -> str:
    """Generate AI response using OpenAI with fallback"""
    if not client:
        # Fallback response when OpenAI is not available
        if context_chunks:
            top_chunk = context_chunks[0]
            return f"Based on the document '{top_chunk.metadata['filename']}', here's what I found:\n\n{top_chunk.content[:500]}..."
        else:
            return "OpenAI API key not configured. Please set OPENAI_API_KEY environment variable."
    
    # Prepare context
    context_text = "\n\n".join([chunk.content for chunk in context_chunks])
    
    # Create prompt
    prompt = f"""You are a helpful AI assistant that answers questions based on the provided document context.

Context from documents:
{context_text}

Question: {query}

Please provide a comprehensive answer based on the context above. If the context doesn't contain enough information to answer the question, please say so and provide what information you can."""

    try:
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You are a helpful AI assistant that answers questions based on document context."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=1000,
            temperature=0.7
        )
        
        ai_response = response.choices[0].message.content
        
        # If AI response is too generic, provide fallback
        if len(ai_response) < 50 or "I don't know" in ai_response.lower() or "I cannot" in ai_response.lower():
            print("AI response seems inadequate, providing fallback")
            if context_chunks:
                top_chunk = context_chunks[0]
                return f"Based on the document '{top_chunk.metadata['filename']}', here's what I found:\n\n{top_chunk.content[:500]}..."
        
        return ai_response
    except Exception as e:
        print(f"Error generating AI response: {str(e)}")
        # Fallback to simple text extraction
        if context_chunks:
            top_chunk = context_chunks[0]
            return f"Based on the document '{top_chunk.metadata['filename']}', here's what I found:\n\n{top_chunk.content[:500]}..."
        else:
            return f"Error generating response: {str(e)}"

# API Routes
@app.get("/")
async def root():
    """Root endpoint - redirect to web UI"""
    return RedirectResponse(url="/web/")

@app.get("/api")
async def api_info():
    """API information endpoint"""
    return {
        "message": settings.app_name,
        "version": settings.app_version,
        "status": "running",
        "deployment": "Vercel Database-Persistent",
        "features": [
            "Document Upload & Processing",
            "Vector Search with OpenAI Embeddings",
            "AI-Powered Responses",
            "Real-time Chat Interface",
            "Document Management",
            "Database Persistence"
        ],
        "endpoints": {
            "web_ui": "/web/",
            "api_docs": "/docs",
            "health": "/health",
            "queries": "/api/v1/queries/",
            "documents": "/api/v1/documents/",
            "upload": "/api/v1/documents/upload",
            "stats": "/api/v1/stats/",
            "debug": "/api/v1/debug"
        }
    }

@app.get("/health")
async def health_check():
    """Health check endpoint"""
    conn = get_db_connection()
    db_status = "connected" if conn else "disconnected"
    if conn:
        conn.close()
    
    return {
        "status": "healthy",
        "version": settings.app_version,
        "deployment": "Vercel Database-Persistent",
        "timestamp": time.strftime("%Y-%m-%d %H:%M:%S"),
        "openai_configured": bool(settings.openai_api_key),
        "database_status": db_status,
        "database_url_configured": bool(settings.database_url)
    }

@app.get("/api/v1/debug")
async def debug_info():
    """Debug endpoint to check system state"""
    conn = get_db_connection()
    
    if conn:
        try:
            cur = conn.cursor()
            cur.execute("SELECT COUNT(*) FROM documents")
            doc_count = cur.fetchone()[0]
            
            cur.execute("SELECT COUNT(*) FROM document_chunks")
            chunk_count = cur.fetchone()[0]
            
            cur.execute("SELECT COUNT(*) FROM queries")
            query_count = cur.fetchone()[0]
            
            cur.close()
            conn.close()
        except Exception as e:
            doc_count = chunk_count = query_count = 0
            print(f"Debug query error: {str(e)}")
    else:
        doc_count = chunk_count = query_count = 0
    
    return {
        "openai_configured": bool(settings.openai_api_key),
        "openai_key_preview": settings.openai_api_key[:10] + "..." if settings.openai_api_key else "Not set",
        "database_configured": bool(settings.database_url),
        "database_url_preview": settings.database_url[:20] + "..." if settings.database_url else "Not set",
        "documents_count": doc_count,
        "chunks_count": chunk_count,
        "queries_count": query_count
    }

@app.post("/api/v1/queries/", response_model=QueryResponse)
async def process_query(request: QueryRequest):
    """Process a query using RAG"""
    start_time = time.time()
    
    try:
        print(f"Processing query: {request.query}")
        
        # Search for similar chunks
        context_chunks = search_similar_chunks(request.query, request.k)
        
        print(f"Found {len(context_chunks)} context chunks")
        
        # Generate response
        if context_chunks:
            response_text = generate_response(request.query, context_chunks)
            context_documents = list(set([chunk.metadata["filename"] for chunk in context_chunks]))
            print(f"Generated response with {len(context_documents)} source documents")
        else:
            response_text = "I couldn't find relevant information in the uploaded documents to answer your question. Please upload some documents first or try rephrasing your question."
            context_documents = []
            print("No context chunks found - returning default response")
        
        processing_time = time.time() - start_time
        
        # Store query in database
        query_id = hashlib.md5(request.query.encode()).hexdigest()[:8]
        conn = get_db_connection()
        if conn:
            try:
                cur = conn.cursor()
                cur.execute("""
                    INSERT INTO queries (query_id, query_text, response, processing_time)
                    VALUES (%s, %s, %s, %s)
                    ON CONFLICT (query_id) DO UPDATE SET
                        query_text = EXCLUDED.query_text,
                        response = EXCLUDED.response,
                        processing_time = EXCLUDED.processing_time
                """, (query_id, request.query, response_text, processing_time))
                conn.commit()
                cur.close()
                conn.close()
            except Exception as e:
                print(f"Error storing query: {str(e)}")
                if conn:
                    conn.close()
        
        return QueryResponse(
            query_id=query_id,
            query=request.query,
            response=response_text,
            context_documents=context_documents,
            context_chunks=context_chunks,
            similarity_scores=[chunk.score for chunk in context_chunks],
            processing_time=processing_time,
            cached=False
        )
        
    except Exception as e:
        print(f"Error processing query: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/v1/documents/upload")
async def upload_document(file: UploadFile = File(...)):
    """Upload and process a document"""
    try:
        # Read file content
        file_content = await file.read()
        
        # Process document
        doc_info = process_document(file_content, file.filename, file.content_type)
        
        return {
            "message": "Document uploaded and processed successfully",
            "document": doc_info
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/v1/documents/", response_model=List[DocumentInfo])
async def get_documents():
    """Get list of uploaded documents"""
    conn = get_db_connection()
    if not conn:
        return []
    
    try:
        cur = conn.cursor(cursor_factory=RealDictCursor)
        cur.execute("SELECT * FROM documents ORDER BY created_at DESC")
        documents = cur.fetchall()
        
        result = []
        for doc in documents:
            result.append(DocumentInfo(
                id=doc['doc_id'],
                filename=doc['filename'],
                file_type=doc['file_type'],
                file_size=doc['file_size'],
                status=doc['status'],
                chunks_count=doc['chunks_count'],
                created_at=doc['created_at'].isoformat() if doc['created_at'] else "",
                metadata=doc['metadata'] if doc['metadata'] else {}
            ))
        
        cur.close()
        conn.close()
        return result
        
    except Exception as e:
        print(f"Error getting documents: {str(e)}")
        if conn:
            conn.close()
        return []

@app.delete("/api/v1/documents/{document_id}")
async def delete_document(document_id: str):
    """Delete a document and its chunks"""
    conn = get_db_connection()
    if not conn:
        raise HTTPException(status_code=500, detail="Database connection failed")
    
    try:
        cur = conn.cursor()
        
        # Delete chunks first
        cur.execute("DELETE FROM document_chunks WHERE document_id = %s", (document_id,))
        
        # Delete document
        cur.execute("DELETE FROM documents WHERE doc_id = %s", (document_id,))
        
        if cur.rowcount == 0:
            raise HTTPException(status_code=404, detail="Document not found")
        
        conn.commit()
        cur.close()
        conn.close()
        
        return {"message": "Document deleted successfully"}
        
    except Exception as e:
        conn.rollback()
        cur.close()
        conn.close()
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/v1/stats/", response_model=StatsResponse)
async def get_stats():
    """Get platform statistics"""
    conn = get_db_connection()
    
    if conn:
        try:
            cur = conn.cursor()
            cur.execute("SELECT COUNT(*) FROM documents")
            doc_count = cur.fetchone()[0]
            
            cur.execute("SELECT COUNT(*) FROM document_chunks")
            chunk_count = cur.fetchone()[0]
            
            cur.execute("SELECT COUNT(*) FROM queries")
            query_count = cur.fetchone()[0]
            
            cur.execute("SELECT AVG(processing_time) FROM queries WHERE processing_time IS NOT NULL")
            avg_time = cur.fetchone()[0] or 0.0
            
            cur.close()
            conn.close()
            
            return StatsResponse(
                total_documents=doc_count,
                total_chunks=chunk_count,
                total_queries=query_count,
                avg_processing_time=float(avg_time)
            )
        except Exception as e:
            print(f"Error getting stats: {str(e)}")
            if conn:
                conn.close()
    
    return StatsResponse(
        total_documents=0,
        total_chunks=0,
        total_queries=0,
        avg_processing_time=0.0
    )

# Web UI (same as before but with database persistence info)
@app.get("/web/", response_class=HTMLResponse)
async def web_ui():
    """Serve the full-featured web UI"""
    html_content = """
    <!DOCTYPE html>
    <html lang="en">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>Document Intelligence Platform</title>
        <style>
            * { margin: 0; padding: 0; box-sizing: border-box; }
            body { 
                font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
                background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                min-height: 100vh;
                display: flex;
                align-items: center;
                justify-content: center;
            }
            .container {
                background: white;
                border-radius: 20px;
                box-shadow: 0 20px 40px rgba(0,0,0,0.1);
                width: 95%;
                max-width: 1200px;
                padding: 30px;
                text-align: center;
            }
            h1 { color: #333; font-size: 2.5rem; margin-bottom: 20px; }
            p { color: #666; font-size: 1.1rem; margin-bottom: 30px; }
            .status { background: #d4edda; color: #155724; padding: 15px; border-radius: 10px; margin-bottom: 20px; }
            .btn { 
                background: linear-gradient(135deg, #667eea, #764ba2);
                color: white; padding: 15px 30px; border: none; border-radius: 25px;
                font-size: 16px; cursor: pointer; margin: 10px; text-decoration: none; display: inline-block;
            }
            .btn:hover { transform: translateY(-2px); }
            .btn-secondary {
                background: linear-gradient(135deg, #28a745, #20c997);
            }
            .btn-danger {
                background: linear-gradient(135deg, #dc3545, #fd7e14);
            }
            .main-content {
                display: grid;
                grid-template-columns: 1fr 300px;
                gap: 30px;
                margin-top: 30px;
            }
            .chat-section {
                text-align: left;
            }
            .sidebar {
                background: #f8f9fa;
                padding: 20px;
                border-radius: 15px;
                text-align: left;
            }
            .chat-container {
                background: #f8f9fa;
                border-radius: 15px;
                padding: 20px;
                margin: 20px 0;
                min-height: 400px;
                max-height: 500px;
                overflow-y: auto;
            }
            .message {
                background: white;
                padding: 15px;
                border-radius: 10px;
                margin: 10px 0;
                box-shadow: 0 2px 5px rgba(0,0,0,0.1);
            }
            .user-message {
                background: #e3f2fd;
                margin-left: 20px;
            }
            .ai-message {
                background: #f1f8e9;
                margin-right: 20px;
            }
            .thinking {
                background: #fff3cd;
                font-style: italic;
            }
            input, textarea {
                width: 100%;
                padding: 15px;
                border: 2px solid #e9ecef;
                border-radius: 25px;
                font-size: 16px;
                margin: 10px 0;
            }
            .upload-area {
                border: 2px dashed #ccc;
                border-radius: 15px;
                padding: 30px;
                margin: 20px 0;
                background: #f9f9f9;
                cursor: pointer;
            }
            .upload-area.dragover {
                border-color: #667eea;
                background: #f0f4ff;
            }
            .document-item {
                background: white;
                padding: 15px;
                border-radius: 8px;
                margin: 10px 0;
                display: flex;
                justify-content: space-between;
                align-items: center;
            }
            .document-info {
                flex: 1;
            }
            .document-actions {
                display: flex;
                gap: 10px;
            }
            .btn-small {
                padding: 8px 15px;
                font-size: 14px;
                margin: 0;
            }
            .stats-grid {
                display: grid;
                grid-template-columns: 1fr 1fr;
                gap: 10px;
                margin: 15px 0;
            }
            .stat-item {
                background: white;
                padding: 10px;
                border-radius: 8px;
                text-align: center;
            }
            .stat-number {
                font-size: 1.5rem;
                font-weight: bold;
                color: #667eea;
            }
            .stat-label {
                font-size: 0.9rem;
                color: #666;
            }
            @media (max-width: 768px) {
                .main-content {
                    grid-template-columns: 1fr;
                }
                .container {
                    padding: 20px;
                }
            }
        </style>
    </head>
    <body>
        <div class="container">
            <h1>🤖 Document Intelligence Platform</h1>
            <p>AI-powered document analysis with RAG capabilities</p>
            <div class="status">✅ Database-persistent version deployed on Vercel!</div>
            <p>Upload documents and ask intelligent questions about their content. Data is now persisted in PostgreSQL.</p>
            
            <div class="main-content">
                <div class="chat-section">
                    <div class="chat-container" id="chatContainer">
                        <div class="message ai-message">
                            👋 Hello! I'm your Document Intelligence assistant. Upload some documents and ask me questions about their content! Data is now persisted in the database.
                        </div>
                    </div>
                    
                    <input type="text" id="queryInput" placeholder="Ask me anything about your uploaded documents..." />
                    <button class="btn" onclick="sendQuery()">Send Query</button>
                </div>
                
                <div class="sidebar">
                    <h3>📁 Document Management</h3>
                    
                    <div class="upload-area" id="uploadArea" onclick="document.getElementById('fileInput').click()">
                        <p>📤 Click to upload or drag & drop files</p>
                        <p style="font-size: 0.9rem; color: #666;">Supports PDF, DOCX, TXT</p>
                        <input type="file" id="fileInput" multiple accept=".pdf,.docx,.txt" style="display: none;" onchange="uploadFiles()">
                    </div>
                    
                    <div id="documentsList">
                        <p style="color: #666; font-style: italic;">No documents uploaded yet</p>
                    </div>
                    
                    <h3>📊 Platform Stats</h3>
                    <div class="stats-grid" id="statsGrid">
                        <div class="stat-item">
                            <div class="stat-number" id="totalDocs">0</div>
                            <div class="stat-label">Documents</div>
                        </div>
                        <div class="stat-item">
                            <div class="stat-number" id="totalChunks">0</div>
                            <div class="stat-label">Chunks</div>
                        </div>
                        <div class="stat-item">
                            <div class="stat-number" id="totalQueries">0</div>
                            <div class="stat-label">Queries</div>
                        </div>
                        <div class="stat-item">
                            <div class="stat-number" id="avgTime">0s</div>
                            <div class="stat-label">Avg Time</div>
                        </div>
                    </div>
                    
                    <div style="margin-top: 20px;">
                        <a href="/api" class="btn btn-small">📚 API Info</a>
                        <a href="/health" class="btn btn-small">❤️ Health</a>
                        <a href="/docs" class="btn btn-small">📖 Docs</a>
                        <a href="/api/v1/debug" class="btn btn-small">🔍 Debug</a>
                    </div>
                </div>
            </div>
        </div>
        
        <script>
            let queryCount = 0;
            
            // Drag and drop functionality
            const uploadArea = document.getElementById('uploadArea');
            
            uploadArea.addEventListener('dragover', (e) => {
                e.preventDefault();
                uploadArea.classList.add('dragover');
            });
            
            uploadArea.addEventListener('dragleave', () => {
                uploadArea.classList.remove('dragover');
            });
            
            uploadArea.addEventListener('drop', (e) => {
                e.preventDefault();
                uploadArea.classList.remove('dragover');
                const files = e.dataTransfer.files;
                handleFiles(files);
            });
            
            async function uploadFiles() {
                const fileInput = document.getElementById('fileInput');
                const files = fileInput.files;
                handleFiles(files);
            }
            
            async function handleFiles(files) {
                for (let file of files) {
                    await uploadFile(file);
                }
                loadDocuments();
                loadStats();
            }
            
            async function uploadFile(file) {
                const formData = new FormData();
                formData.append('file', file);
                
                try {
                    const response = await fetch('/api/v1/documents/upload', {
                        method: 'POST',
                        body: formData
                    });
                    
                    const result = await response.json();
                    
                    if (response.ok) {
                        addMessage(`📄 Document "${file.name}" uploaded and processed successfully!`, 'ai-message');
                    } else {
                        addMessage(`❌ Error uploading "${file.name}": ${result.detail}`, 'ai-message');
                    }
                } catch (error) {
                    addMessage(`❌ Error uploading "${file.name}": ${error.message}`, 'ai-message');
                }
            }
            
            async function sendQuery() {
                const input = document.getElementById('queryInput');
                const query = input.value.trim();
                
                if (!query) return;
                
                queryCount++;
                
                // Add user message
                addMessage(query, 'user-message');
                input.value = '';
                
                // Show thinking indicator
                const thinkingId = addMessage('🤔 Thinking...', 'ai-message thinking');
                
                try {
                    const response = await fetch('/api/v1/queries/', {
                        method: 'POST',
                        headers: {
                            'Content-Type': 'application/json',
                        },
                        body: JSON.stringify({
                            query: query,
                            k: 5,
                            use_cache: false,
                            rerank_context: true
                        })
                    });
                    
                    const data = await response.json();
                    
                    // Remove thinking message
                    removeMessage(thinkingId);
                    
                    // Add AI response
                    addMessage(data.response, 'ai-message');
                    
                    // Show context if available
                    if (data.context_documents && data.context_documents.length > 0) {
                        addMessage(`📚 Sources: ${data.context_documents.join(', ')}`, 'ai-message');
                    }
                    
                } catch (error) {
                    removeMessage(thinkingId);
                    addMessage(`❌ Error: ${error.message}`, 'ai-message');
                }
            }
            
            function addMessage(content, className) {
                const chatContainer = document.getElementById('chatContainer');
                const messageDiv = document.createElement('div');
                messageDiv.className = `message ${className}`;
                messageDiv.innerHTML = `<strong>${className.includes('user') ? 'You' : 'AI'}:</strong> ${content}`;
                chatContainer.appendChild(messageDiv);
                chatContainer.scrollTop = chatContainer.scrollHeight;
                return messageDiv;
            }
            
            function removeMessage(messageElement) {
                if (messageElement && messageElement.parentNode) {
                    messageElement.parentNode.removeChild(messageElement);
                }
            }
            
            async function loadDocuments() {
                try {
                    const response = await fetch('/api/v1/documents/');
                    const documents = await response.json();
                    
                    const documentsList = document.getElementById('documentsList');
                    
                    if (documents.length === 0) {
                        documentsList.innerHTML = '<p style="color: #666; font-style: italic;">No documents uploaded yet</p>';
                        return;
                    }
                    
                    documentsList.innerHTML = documents.map(doc => `
                        <div class="document-item">
                            <div class="document-info">
                                <strong>📄 ${doc.filename}</strong><br>
                                <small>${(doc.file_size / 1024).toFixed(0)} KB • ${doc.chunks_count} chunks</small>
                            </div>
                            <div class="document-actions">
                                <button class="btn btn-small btn-secondary" onclick="viewDocument('${doc.id}')">View</button>
                                <button class="btn btn-small btn-danger" onclick="deleteDocument('${doc.id}')">Delete</button>
                            </div>
                        </div>
                    `).join('');
                    
                } catch (error) {
                    console.error('Error loading documents:', error);
                }
            }
            
            async function loadStats() {
                try {
                    const response = await fetch('/api/v1/stats/');
                    const stats = await response.json();
                    
                    document.getElementById('totalDocs').textContent = stats.total_documents;
                    document.getElementById('totalChunks').textContent = stats.total_chunks;
                    document.getElementById('totalQueries').textContent = stats.total_queries;
                    document.getElementById('avgTime').textContent = `${stats.avg_processing_time.toFixed(2)}s`;
                    
                } catch (error) {
                    console.error('Error loading stats:', error);
                }
            }
            
            async function deleteDocument(docId) {
                if (!confirm('Are you sure you want to delete this document?')) return;
                
                try {
                    const response = await fetch(`/api/v1/documents/${docId}`, {
                        method: 'DELETE'
                    });
                    
                    if (response.ok) {
                        addMessage(`🗑️ Document deleted successfully`, 'ai-message');
                        loadDocuments();
                        loadStats();
                    } else {
                        const result = await response.json();
                        addMessage(`❌ Error deleting document: ${result.detail}`, 'ai-message');
                    }
                } catch (error) {
                    addMessage(`❌ Error deleting document: ${error.message}`, 'ai-message');
                }
            }
            
            function viewDocument(docId) {
                addMessage(`📄 Viewing document ${docId}`, 'ai-message');
            }
            
            // Handle Enter key
            document.getElementById('queryInput').addEventListener('keypress', function(e) {
                if (e.key === 'Enter') {
                    sendQuery();
                }
            });
            
            // Load initial data
            loadDocuments();
            loadStats();
        </script>
    </body>
    </html>
    """
    return HTMLResponse(content=html_content)

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
