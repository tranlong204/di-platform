"""
Document Intelligence Platform - Minimal Vercel Version
"""

from fastapi import FastAPI, UploadFile, File, Form, Depends, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import RedirectResponse
from fastapi.staticfiles import StaticFiles
from sqlalchemy.orm import Session
from datetime import datetime
import os
import json
import PyPDF2
from docx import Document
from io import BytesIO
import zipfile
import xml.etree.ElementTree as ET
from supabase import create_client, Client
from pydantic_settings import BaseSettings
from openai import OpenAI

# Configuration
class Settings(BaseSettings):
    supabase_url: str
    supabase_key: str
    supabase_service_key: str = ""
    openai_api_key: str = ""
    api_base_url: str = ""
    
    class Config:
        env_file = ".env"

settings = Settings()

# Initialize Supabase client
supabase_server_key = settings.supabase_service_key or settings.supabase_key
supabase: Client = create_client(settings.supabase_url, supabase_server_key)

# Initialize OpenAI client
openai_client = OpenAI(api_key=settings.openai_api_key) if settings.openai_api_key else None

# Simple in-memory storage for demo (in production, use Supabase)
documents_storage = []
queries_storage = []

# Database dependency (simplified for Vercel)
def get_db():
    # In a real implementation, this would return a database session
    # For now, we'll use in-memory storage
    return None


def extract_text_from_document(content: bytes, file_extension: str) -> str:
    """Extract text from uploaded document"""
    try:
        if file_extension == ".pdf":
            # Extract text from PDF
            pdf_reader = PyPDF2.PdfReader(BytesIO(content))
            pages_text = []
            for page in pdf_reader.pages:
                # Some PDFs/pages return None for extract_text(); skip safely.
                page_text = page.extract_text() or ""
                if page_text.strip():
                    pages_text.append(page_text)
            extracted = "\n".join(pages_text).strip()
            if extracted:
                return extracted

            # Fallback 1: pdfplumber often extracts text from PDFs that PyPDF2 cannot.
            try:
                import pdfplumber
                fallback_pages = []
                with pdfplumber.open(BytesIO(content)) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text() or ""
                        if page_text.strip():
                            fallback_pages.append(page_text)
                plumber_text = "\n".join(fallback_pages).strip()
                if plumber_text:
                    return plumber_text
            except Exception as fallback_error:
                print(f"pdfplumber fallback failed: {fallback_error}")

            # Fallback 2: AWS Textract OCR — handles image-based / scanned PDFs.
            # Split into single pages so we can use the sync DetectDocumentText API.
            try:
                import boto3
                textract = boto3.client("textract", region_name=os.getenv("AWS_REGION", "us-west-1"))
                pdf_reader = PyPDF2.PdfReader(BytesIO(content))
                all_text = []
                for page_num in range(len(pdf_reader.pages)):
                    writer = PyPDF2.PdfWriter()
                    writer.add_page(pdf_reader.pages[page_num])
                    page_buf = BytesIO()
                    writer.write(page_buf)
                    page_bytes = page_buf.getvalue()
                    response = textract.detect_document_text(Document={"Bytes": page_bytes})
                    lines = [
                        block["Text"]
                        for block in response.get("Blocks", [])
                        if block.get("BlockType") == "LINE"
                    ]
                    if lines:
                        all_text.append("\n".join(lines))
                textract_text = "\n\n".join(all_text).strip()
                if textract_text:
                    print(f"Textract extracted {len(textract_text)} chars from {len(pdf_reader.pages)} pages")
                    return textract_text
            except Exception as textract_error:
                print(f"Textract fallback failed: {textract_error}")

            return ""
        
        elif file_extension == ".docx":
            # Extract text from DOCX
            try:
                doc = Document(BytesIO(content))
                text = ""
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                extracted = text.strip()
                if extracted:
                    return extracted
            except Exception as docx_error:
                print(f"python-docx extraction failed: {docx_error}")

            # Fallback: parse DOCX XML directly (DOCX is a zip archive).
            # Read all Word XML parts (document, headers, footers, notes, etc.)
            # to handle files where body text is not only in word/document.xml.
            try:
                with zipfile.ZipFile(BytesIO(content)) as zf:
                    xml_files = sorted(
                        name for name in zf.namelist()
                        if name.startswith("word/") and name.endswith(".xml")
                    )
                    text_nodes = []
                    for xml_name in xml_files:
                        try:
                            xml_data = zf.read(xml_name)
                            root = ET.fromstring(xml_data)
                            for node in root.iter():
                                # Collect all tags ending with 't' (text runs), plus instrText.
                                tag = node.tag
                                if not isinstance(tag, str):
                                    continue
                                if tag.endswith("}t") or tag.endswith("}instrText"):
                                    if node.text and node.text.strip():
                                        text_nodes.append(node.text.strip())
                        except Exception as xml_part_error:
                            print(f"DOCX XML part parse failed ({xml_name}): {xml_part_error}")
                            continue
                extracted = "\n".join(text_nodes).strip()
                if extracted:
                    return extracted
                return ""
            except Exception as fallback_error:
                print(f"DOCX XML fallback failed: {fallback_error}")
                return ""
        
        elif file_extension == ".txt":
            # Extract text from TXT
            return content.decode('utf-8').strip()
        
        else:
            return ""
    
    except Exception as e:
        print(f"Error extracting text: {e}")
        return ""


def generate_chatgpt_response(query: str, document_content: str, filename: str) -> str:
    """Generate response using ChatGPT API"""
    if not openai_client:
        return "ChatGPT API not configured. Please set OPENAI_API_KEY environment variable."
    
    try:
        # Create a context-aware prompt
        system_prompt = f"""You are a helpful document intelligence assistant. You have access to the following document content and should answer questions based on it.

Document: {filename}
Content: {document_content[:4000]}  # Limit content to avoid token limits

Please provide a helpful, accurate response based on the document content. If the question cannot be answered from the document, say so clearly."""

        response = openai_client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": query}
            ],
            max_tokens=1000,
            temperature=0.7
        )
        
        return response.choices[0].message.content.strip()
    
    except Exception as e:
        print(f"ChatGPT API error: {e}")
        return f"Error generating response: {str(e)}"


def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200) -> list:
    """Split text into overlapping chunks"""
    if not text:
        return []
    
    chunks = []
    start = 0
    
    while start < len(text):
        end = start + chunk_size
        
        # Try to break at sentence boundary
        if end < len(text):
            # Look for sentence endings within the last 100 characters
            for i in range(end, max(start + chunk_size - 100, start), -1):
                if text[i] in '.!?':
                    end = i + 1
                    break
        
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        
        # Move start position with overlap
        start = end - overlap
        if start >= len(text):
            break
    
    return chunks


# Create FastAPI application
app = FastAPI(
    title="Document Intelligence Platform",
    version="1.0.0",
    description="AI-powered document assistant with RAG capabilities"
)

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Mount static files (web UI)
web_dir = "web"
if os.path.exists(web_dir):
    app.mount("/web", StaticFiles(directory=web_dir, html=True), name="web")


@app.get("/")
async def root():
    """Root endpoint - redirect to web UI"""
    return RedirectResponse(url="/web/")


@app.get("/api/v1/stats")
async def get_stats():
    """Get system statistics"""
    processed_docs = sum(1 for doc in documents_storage if doc.get("processed", False))
    total_chunks = sum(doc.get("chunks", 0) for doc in documents_storage)
    avg_processing_time = sum(q.get("processing_time", 0) for q in queries_storage) / len(queries_storage) if queries_storage else 0
    
    return {
        "documents": {"total": len(documents_storage), "processed": processed_docs, "pending": len(documents_storage) - processed_docs},
        "queries": {"total": len(queries_storage), "average_processing_time": avg_processing_time},
        "agents": {"total_executions": 0, "successful_executions": 0, "success_rate": 0.0},
        "vector_store": {"total_vectors": total_chunks, "dimension": 0, "index_type": "none", "documents_count": len(documents_storage)},
        "cache": {"total_keys": 0, "memory_usage": "0B", "hit_rate": 0.0, "connected_clients": 0}
    }


@app.get("/api/v1/documents/")
async def get_documents():
    """Get all documents"""
    try:
        # Try to fetch from Supabase first
        supabase_response = supabase.table("documents").select("*").execute()
        if supabase_response.data:
            return {"documents": supabase_response.data, "total": len(supabase_response.data)}
    except Exception as e:
        print(f"Supabase fetch failed: {e}")
    
    # Fallback to memory storage
    return {"documents": documents_storage, "total": len(documents_storage)}


@app.get("/api/v1/queries/history")
async def get_query_history():
    """Get query history"""
    try:
        # Try to fetch from Supabase first
        supabase_response = supabase.table("queries").select("*").order("created_at", desc=True).execute()
        if supabase_response.data:
            return {"queries": supabase_response.data, "total": len(supabase_response.data)}
    except Exception as e:
        print(f"Supabase fetch failed: {e}")
    
    # Fallback to memory storage
    return {"queries": queries_storage, "total": len(queries_storage)}


@app.get("/api/v1/documents/{document_id}/chunks")
async def get_document_chunks(document_id: int):
    """Get chunks for a specific document"""
    try:
        # Fetch chunks from Supabase
        supabase_response = supabase.table("document_chunks").select("*").eq("document_id", document_id).order("chunk_index").execute()
        if supabase_response.data:
            return {"chunks": supabase_response.data, "total": len(supabase_response.data), "document_id": document_id}
    except Exception as e:
        print(f"Supabase fetch failed: {e}")
    
    return {"chunks": [], "total": 0, "document_id": document_id}


@app.get("/api/v1/documents/{document_id}")
async def get_document(document_id: int):
    """Get a specific document by ID"""
    try:
        # Fetch document from Supabase
        supabase_response = supabase.table("documents").select("*").eq("id", document_id).execute()
        if supabase_response.data:
            doc = supabase_response.data[0]
            return {
                "id": doc["id"],
                "filename": doc["filename"],
                "file_type": doc["file_type"],
                "file_size": doc["file_size"],
                "chunks_count": doc.get("metadata_json", {}).get("chunks", 0),
                "processed": doc["processed"],
                "created_at": doc["created_at"]
            }
    except Exception as e:
        print(f"Supabase fetch failed: {e}")
    
    return {"error": f"Document {document_id} not found"}


@app.delete("/api/v1/documents/{document_id}")
async def delete_document(document_id: int):
    """Delete a document and its chunks"""
    try:
        # First, delete all chunks associated with this document
        try:
            supabase.table("document_chunks").delete().eq("document_id", document_id).execute()
            print(f"Deleted chunks for document {document_id}")
        except Exception as chunk_error:
            print(f"Error deleting chunks: {chunk_error}")
        
        # Then delete the document itself
        try:
            supabase_response = supabase.table("documents").delete().eq("id", document_id).execute()
            if supabase_response.data:
                return {
                    "message": f"Document {document_id} deleted successfully",
                    "deleted_chunks": True,
                    "document_id": document_id
                }
            else:
                return {"error": f"Document {document_id} not found"}
        except Exception as doc_error:
            return {"error": f"Failed to delete document: {str(doc_error)}"}
    
    except Exception as e:
        return {"error": f"Document deletion failed: {str(e)}"}


@app.post("/api/v1/queries/")
async def process_query(query_data: dict):
    """Process a query/question"""
    try:
        query_text = query_data.get("query", "")
        if not query_text:
            return {"error": "No query provided"}
        
        # Create query record
        query_id = len(queries_storage) + 1
        query_record = {
            "id": query_id,
            "query": query_text,
            "response": "",
            "created_at": datetime.now().isoformat(),
            "processing_time": 0
        }
        
        # Generate response using ChatGPT API
        response = "I'm here to help answer questions about your uploaded documents."
        
        # Search through uploaded documents for relevant content
        # First try to get documents from Supabase
        documents_to_search = []
        try:
            supabase_response = supabase.table("documents").select("*").execute()
            if supabase_response.data:
                documents_to_search = supabase_response.data
        except Exception as e:
            print(f"Supabase fetch failed: {e}")
        
        # Fallback to memory storage if Supabase fails
        if not documents_to_search:
            documents_to_search = documents_storage
        
        if documents_to_search:
            # Use ChatGPT to generate intelligent responses
            for doc in documents_to_search:
                # Get content from either Supabase metadata or memory storage
                full_content = ""
                if doc.get("metadata_json") and doc["metadata_json"].get("full_content"):
                    full_content = doc["metadata_json"]["full_content"]
                elif doc.get("full_content"):
                    full_content = doc["full_content"]
                
                if full_content:
                    # Generate ChatGPT response
                    response = generate_chatgpt_response(query_text, full_content, doc['filename'])
                    break
        else:
            response = "I don't see any uploaded documents. Please upload a document first, then ask questions about its content."
        
        query_record["response"] = response
        query_record["processing_time"] = 0.5  # Simulated processing time
        
        # Store the query in Supabase
        try:
            # Save to Supabase queries table
            supabase_response = supabase.table("queries").insert({
                "query_text": query_text,
                "response_text": response,
                "query_hash": f"hash_{len(query_text)}_{hash(query_text)}",
                "processing_time": query_record["processing_time"]
            }).execute()
            
            # Also store in memory for immediate access
            queries_storage.append(query_record)
            
            return {
                "query_id": query_id,
                "response": response,
                "processing_time": query_record["processing_time"],
                "supabase_id": supabase_response.data[0]["id"] if supabase_response.data else None
            }
        except Exception as db_error:
            # Fallback to memory storage if Supabase fails
            queries_storage.append(query_record)
            return {
                "query_id": query_id,
                "response": response,
                "processing_time": query_record["processing_time"],
                "warning": f"Database save failed: {str(db_error)}"
            }
        
    except Exception as e:
        return {"error": f"Query processing failed: {str(e)}"}


@app.post("/api/v1/documents/upload")
async def upload_document(file: UploadFile = File(...)):
    """Upload a document"""
    try:
        # Check file size (2MB limit)
        content = await file.read()
        if len(content) > 2 * 1024 * 1024:  # 2MB
            return {"error": "File too large. Maximum size is 2MB."}
        
        # Check file type
        allowed_types = [".pdf", ".docx", ".txt"]
        file_extension = os.path.splitext(file.filename)[1].lower()
        if file_extension not in allowed_types:
            return {"error": f"File type not supported. Allowed types: {', '.join(allowed_types)}"}
        
        # Create document record
        document_id = len(documents_storage) + 1
        
        # Extract text from document
        extracted_text = extract_text_from_document(content, file_extension)
        
        # Create chunks from the extracted text
        chunks = chunk_text(extracted_text) if extracted_text else []

        # Log extraction result for diagnosis but do not hard-block upload.
        # Documents with no extractable text are stored with processed=False.
        if not extracted_text or not chunks:
            print(
                f"[WARN] No text extracted from {file.filename} "
                f"(ext={file_extension}, size={len(content)}). "
                f"Document will be stored as unprocessed."
            )
        
        document_record = {
            "id": document_id,
            "filename": file.filename,
            "size": len(content),
            "type": file_extension,
            "uploaded_at": datetime.now().isoformat(),
            "processed": True if extracted_text else False,
            "chunks": len(chunks),
            "content": extracted_text[:1000] if extracted_text else "",  # Store first 1000 chars for demo
            "full_content": extracted_text  # Store full content for querying
        }
        
        # Store the document in Supabase.
        # This endpoint must persist to DB; do not silently degrade to memory-only mode.
        try:
            # Save to Supabase documents table
            supabase_response = supabase.table("documents").insert({
                "filename": file.filename,
                "file_path": f"/uploads/{file.filename}",  # Virtual path
                "file_type": file_extension,
                "file_size": len(content),
                "content_hash": f"hash_{len(content)}_{file.filename}",  # Simple hash
                "title": file.filename,
                "summary": extracted_text[:500] if extracted_text else "",
                "metadata_json": {"full_content": extracted_text, "chunks": len(chunks)},
                "processed": True if extracted_text else False
            }).execute()
            
            document_supabase_id = supabase_response.data[0]["id"] if supabase_response.data else None
            
            # Save chunks to document_chunks table
            if chunks and document_supabase_id:
                chunk_records = []
                for i, chunk in enumerate(chunks):
                    chunk_record = {
                        "document_id": document_supabase_id,
                        "chunk_index": i,
                        "content": chunk,
                        "content_hash": f"chunk_{i}_{hash(chunk)}",
                        "token_count": len(chunk.split()),  # Approximate token count
                        "metadata_json": {"chunk_size": len(chunk), "overlap": 200 if i > 0 else 0}
                    }
                    chunk_records.append(chunk_record)
                
                # Insert all chunks at once
                if chunk_records:
                    supabase.table("document_chunks").insert(chunk_records).execute()
            
            # Also store in memory for immediate access
            documents_storage.append(document_record)
            
            return {
                "message": f"Document '{file.filename}' uploaded successfully",
                "document": document_record,
                "supabase_id": document_supabase_id,
                "chunks_created": len(chunks)
            }
        except Exception as db_error:
            raise HTTPException(
                status_code=500,
                detail=f"Upload failed to persist in Supabase: {str(db_error)}"
            )
        
    except HTTPException:
        raise
    except Exception as e:
        return {"error": f"Upload failed: {str(e)}"}


@app.get("/api")
async def api_info():
    """API information endpoint"""
    return {
        "message": "Document Intelligence Platform API",
        "version": "1.0.0",
        "status": "running",
        "endpoints": {
            "web_ui": "/web/",
            "api_docs": "/docs",
            "health": "/health"
        }
    }


@app.post("/api/v1/debug/extract")
async def debug_extract(file: UploadFile = File(...)):
    """Debug endpoint: returns extraction result without saving anything."""
    content = await file.read()
    file_extension = os.path.splitext(file.filename)[1].lower()
    diagnostics = {
        "filename": file.filename,
        "extension": file_extension,
        "size_bytes": len(content),
        "pypdf2_result": None,
        "pdfplumber_result": None,
        "docx_paragraphs": None,
        "docx_xml_nodes": None,
        "final_text_preview": None,
        "chunks_count": 0,
        "errors": []
    }
    if file_extension == ".pdf":
        try:
            pdf_reader = PyPDF2.PdfReader(BytesIO(content))
            pages_text = []
            for i, page in enumerate(pdf_reader.pages):
                t = page.extract_text() or ""
                pages_text.append(t)
            full = "\n".join(pages_text).strip()
            diagnostics["pypdf2_result"] = {
                "pages": len(pdf_reader.pages),
                "chars": len(full),
                "preview": full[:300]
            }
        except Exception as e:
            diagnostics["errors"].append(f"PyPDF2: {e}")
        try:
            import pdfplumber
            fallback_pages = []
            with pdfplumber.open(BytesIO(content)) as pdf:
                for page in pdf.pages:
                    t = page.extract_text() or ""
                    fallback_pages.append(t)
            full = "\n".join(fallback_pages).strip()
            diagnostics["pdfplumber_result"] = {
                "pages": len(fallback_pages),
                "chars": len(full),
                "preview": full[:300]
            }
        except Exception as e:
            diagnostics["errors"].append(f"pdfplumber: {e}")
    elif file_extension == ".docx":
        try:
            doc = Document(BytesIO(content))
            paras = [p.text for p in doc.paragraphs]
            full = "\n".join(paras).strip()
            diagnostics["docx_paragraphs"] = {"count": len(paras), "chars": len(full), "preview": full[:300]}
        except Exception as e:
            diagnostics["errors"].append(f"python-docx: {e}")
        try:
            with zipfile.ZipFile(BytesIO(content)) as zf:
                xml_files = sorted(n for n in zf.namelist() if n.startswith("word/") and n.endswith(".xml"))
                nodes = []
                for xml_name in xml_files:
                    xml_data = zf.read(xml_name)
                    root = ET.fromstring(xml_data)
                    for node in root.iter():
                        tag = node.tag
                        if isinstance(tag, str) and (tag.endswith("}t") or tag.endswith("}instrText")):
                            if node.text and node.text.strip():
                                nodes.append(node.text.strip())
                full = "\n".join(nodes).strip()
                diagnostics["docx_xml_nodes"] = {"xml_files": xml_files, "nodes_found": len(nodes), "chars": len(full), "preview": full[:300]}
        except Exception as e:
            diagnostics["errors"].append(f"docx-xml: {e}")
    extracted = extract_text_from_document(content, file_extension)
    diagnostics["final_text_preview"] = extracted[:300] if extracted else ""
    diagnostics["chunks_count"] = len(chunk_text(extracted)) if extracted else 0
    return diagnostics


@app.get("/api/config")
async def api_config():
    """Return runtime frontend configuration."""
    api_base = settings.api_base_url.strip() if settings.api_base_url else "/api/v1"
    return {"api_base": api_base}


@app.get("/health")
async def health_check():
    """Health check endpoint"""
    return {"status": "healthy", "timestamp": "2024-01-01T00:00:00Z"}


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(
        "main:app",
        host="0.0.0.0",
        port=8000,
        reload=False,
        log_level="info"
    )